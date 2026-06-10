"""
Microsoft Word 2007 Format
"""

import random

from docx import Document
from docx.shared import Inches, Mm
from htmldocx import HtmlToDocx
from logzero import logger
import re
import config

OPTION_INDENT = Inches(0.25)
MATCHING_BATCH_SIZE = 5
CHOICE_QUESTION_TYPES = (
    "multiple_choice_question",
    "multiple_answers_question",
    "true_false_question",
)


def _set_option_indent(paragraph):
    paragraph.paragraph_format.left_indent = OPTION_INDENT


def _add_option_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    _set_option_indent(paragraph)
    return paragraph


def _assessment_title(assessment):
    title = assessment['metadata']['title']
    form_label = assessment['metadata'].get('form_label')
    if form_label:
        return f"{title} — {form_label}"
    return title


def _option_label(index):
    """Return option letter for index: 0 -> A, 25 -> Z, 26 -> AA."""
    label = ""
    n = index
    while True:
        label = chr(ord("A") + n % 26) + label
        n = n // 26 - 1
        if n < 0:
            break
    return label


def _is_generic_title(title):
    return not title or title.strip().lower() == "question"


def _write_text(doc, html_parser, text):
    if not text:
        return
    if "<" in text:
        html_parser.add_html_to_document(re.sub("</*tbody>", "", text), doc)
    else:
        doc.add_paragraph(text)


def _split_even_batches(items, max_size=MATCHING_BATCH_SIZE):
    """Split items into as-even-as-possible batches of at most max_size."""
    count = len(items)
    if count == 0:
        return []
    batch_count = (count + max_size - 1) // max_size
    base_size = count // batch_count
    extra = count % batch_count
    batches = []
    start = 0
    for index in range(batch_count):
        size = base_size + (1 if index < extra else 0)
        batches.append(items[start:start + size])
        start += size
    return batches


def _matching_bank_for_batch(batch_rows):
    """Word bank terms for a matching batch (correct answers, unique)."""
    seen = set()
    bank = []
    for row in batch_rows:
        for option in row.get("options", []):
            option_id = option.get("id")
            if option.get("correct") and option.get("display") and option.get("text"):
                if option_id not in seen:
                    seen.add(option_id)
                    bank.append(option)
    return bank[:MATCHING_BATCH_SIZE]


def _full_matching_bank(batch_rows):
    if not batch_rows:
        return []
    return [
        option
        for option in batch_rows[0].get("options", [])
        if option.get("display") and option.get("text")
    ]


def prepare_matching_batches(question, split_matching, rng):
    """Build matching groups with a randomized word bank for each group."""
    rows = question.get("answer", [])
    row_batches = (
        _split_even_batches(rows)
        if split_matching and len(rows) > MATCHING_BATCH_SIZE
        else [rows]
    )
    batches = []
    for batch_rows in row_batches:
        if split_matching:
            bank = _matching_bank_for_batch(batch_rows)
        else:
            bank = _full_matching_bank(batch_rows)
        bank = list(bank)
        rng.shuffle(bank)
        batches.append({"rows": batch_rows, "bank": bank})
    return batches


def _prepare_assessment_matching(assessment, split_matching, rng):
    for question in assessment.get("question", []):
        if question.get("question_type") == "matching_question":
            question["_matching_batches"] = prepare_matching_batches(
                question, split_matching, rng
            )


def _write_matching_table(doc, batch, question_num):
    batch_rows = batch["rows"]
    bank = batch["bank"]

    doc.add_paragraph().add_run("Matching").bold = True

    left_lines = []
    for row in batch_rows:
        if row.get("text"):
            left_lines.append(f"{question_num}. {row['text']}")
            question_num += 1

    right_lines = []
    for oindex, option in enumerate(bank):
        right_lines.append(f"{_option_label(oindex)}. {option['text']}")

    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "\n".join(left_lines)
    table.cell(0, 1).text = "\n".join(right_lines)
    return question_num


def _write_matching(doc, question, question_num, html_parser, split_matching=False):
    if question.get("text"):
        _write_text(doc, html_parser, question["text"])

    batches = question.get("_matching_batches", [])

    for batch_index, batch in enumerate(batches):
        if batch_index > 0:
            doc.add_paragraph()
        question_num = _write_matching_table(doc, batch, question_num)

    return question_num


def _write_numbered_prompt(doc, question_num, prompt, html_parser):
    if not prompt:
        doc.add_paragraph(f"{question_num}.")
        return
    if "<" in prompt:
        html_parser.add_html_to_document(
            f"<p><strong>{question_num}.</strong> {re.sub('</*tbody>', '', prompt)}</p>",
            doc,
        )
    else:
        doc.add_paragraph(f"{question_num}. {prompt}")


def _write_choice_question(doc, question, question_num, html_parser):
    _write_numbered_prompt(doc, question_num, question.get("text") or "", html_parser)
    question_num += 1

    for index, answer in enumerate(question.get("answer", [])):
        if not answer.get("display"):
            _add_option_paragraph(doc, config.blanks_replace_str * config.blanks_answer_n)
            continue
        label = _option_label(index)
        if answer.get("text"):
            option_text = answer["text"]
            if "<" in option_text:
                html_parser.add_html_to_document(
                    f'<p style="margin-left: 0.25in">{label}. {option_text}</p>',
                    doc,
                )
            else:
                _add_option_paragraph(doc, f"{label}. {option_text}")
        elif answer.get("image"):
            _add_option_paragraph(doc, f"{label}.")
            for img in answer["image"]:
                doc.add_picture(img["href"].replace("%20", " "), height=Mm(10))

    return question_num


def _write_generic_question(doc, question, question_num, html_parser):
    if not _is_generic_title(question.get("title")):
        doc.add_heading(question["title"], level=2)

    if question.get("image"):
        for img in question["image"]:
            doc.add_picture(img["href"].replace("%20", " "), width=Mm(100))

    prompt = question.get("text") or ""
    if prompt:
        _write_numbered_prompt(doc, question_num, prompt, html_parser)
        question_num += 1

    qtype = question.get("question_type")
    if qtype == "multiple_dropdowns_question":
        for aindex, group in enumerate(question.get("answer", [])):
            options = []
            for option in group.get("options", []):
                if option.get("display"):
                    options.append(option.get("text") or "---")
            doc.add_paragraph(f"{aindex + 1}: " + ", ".join(map(str, options)))
    elif qtype == "calculated_question":
        if config.calculated_display_var_set_in_text:
            doc.add_paragraph(config.blanks_replace_str * config.blanks_answer_n)
        else:
            for index, answer in enumerate(question.get("answer", [])):
                if answer.get("display") and answer.get("text"):
                    doc.add_paragraph(
                        f"{_option_label(index)}. {answer['text']}: "
                        + config.blanks_replace_str * 20
                    )
    elif question.get("answer"):
        for index, answer in enumerate(question["answer"]):
            if answer.get("display") and answer.get("text"):
                doc.add_paragraph(f"{_option_label(index)}. {answer['text']}")
            elif not answer.get("display"):
                doc.add_paragraph(config.blanks_replace_str * config.blanks_answer_n)

    return question_num


def _write_section_header(doc, question):
    qtype = question.get("question_type")
    if qtype == "matching_question":
        return
    if qtype in CHOICE_QUESTION_TYPES and question.get("section_id"):
        doc.add_paragraph().add_run("Multiple Choice").bold = True


def _write_question(doc, question, question_num, html_parser, split_matching=False):
    qtype = question.get("question_type")

    if qtype == "matching_question":
        return _write_matching(doc, question, question_num, html_parser, split_matching)
    if qtype in CHOICE_QUESTION_TYPES:
        return _write_choice_question(doc, question, question_num, html_parser)
    return _write_generic_question(doc, question, question_num, html_parser)


def _matching_key_entries(batch, question_num):
    entries = []
    batch_rows = batch["rows"]
    bank = batch["bank"]
    label_map = {option["id"]: _option_label(index) for index, option in enumerate(bank)}

    for row in batch_rows:
        if not row.get("text"):
            continue
        letter = None
        for option in row.get("options", []):
            if option.get("correct"):
                letter = label_map.get(option["id"])
                break
        if letter:
            entries.append(f"{question_num}. {letter}")
        question_num += 1
    return entries, question_num


def _choice_key_entry(question, question_num):
    labels = []
    for index, answer in enumerate(question.get("answer", [])):
        if answer.get("correct") and answer.get("display"):
            labels.append(_option_label(index))
    if not labels:
        return None
    return f"{question_num}. {', '.join(labels)}"


def _text_key_entry(question, question_num):
    texts = [
        answer["text"]
        for answer in question.get("answer", [])
        if answer.get("correct") and answer.get("text")
    ]
    if not texts:
        return None
    return f"{question_num}. {', '.join(texts)}"


def _build_answer_key_lines(assessment):
    lines = []
    question_num = 1

    for question in assessment.get("question", []):
        qtype = question.get("question_type")
        if qtype == "matching_question":
            for batch in question.get("_matching_batches", []):
                entries, question_num = _matching_key_entries(batch, question_num)
                lines.extend(entries)
        elif qtype in CHOICE_QUESTION_TYPES:
            entry = _choice_key_entry(question, question_num)
            if entry:
                lines.append(entry)
            question_num += 1
        else:
            entry = _text_key_entry(question, question_num)
            if entry:
                lines.append(entry)
            question_num += 1

    return lines


def _write_answer_key(doc, assessment):
    lines = _build_answer_key_lines(assessment)
    if not lines:
        return

    doc.add_page_break()
    doc.add_heading("Answer Key", level=0)
    doc.add_heading(_assessment_title(assessment), level=1)
    for line in lines:
        doc.add_paragraph(line)


def _write_assessment_content(doc, assessment, html_parser, split_matching=False, output_key=False, rng=None):
    if rng is None:
        rng = random.Random()
    _prepare_assessment_matching(assessment, split_matching, rng)

    doc.add_heading(_assessment_title(assessment), 0)

    description = assessment['metadata'].get('description') or ''
    logger.info("Writing assessment: " + _assessment_title(assessment))
    logger.info("with description: " + description)
    if description:
        _write_text(doc, html_parser, description)

    question_num = 1
    current_section_id = object()
    for index, question in enumerate(assessment['question']):
        section_id = question.get("section_id")
        if section_id != current_section_id:
            if index > 0:
                doc.add_paragraph()
            _write_section_header(doc, question)
            current_section_id = section_id
        elif index > 0:
            doc.add_paragraph()
        question_num = _write_question(doc, question, question_num, html_parser, split_matching)

    if output_key:
        _write_answer_key(doc, assessment)


def write_file(data, outfile, split_matching=False, output_key=False, rng=None):
    doc = Document()
    doc = setup_a4(doc)
    doc = setup_metadata(doc)

    html_parser = HtmlToDocx()

    for assessment in data['assessment']:
        _write_assessment_content(
            doc, assessment, html_parser, split_matching, output_key=output_key, rng=rng
        )

    doc.save(outfile)


def write_forms(forms, outfile, combined=False, split_matching=False, output_key=False, seed=None):
    doc = Document()
    doc = setup_a4(doc)
    doc = setup_metadata(doc)

    html_parser = HtmlToDocx()

    for form_index, form in enumerate(forms):
        rng = random.Random(seed + form_index) if seed is not None else random.Random()
        for assessment in form['assessment']:
            _write_assessment_content(
                doc,
                assessment,
                html_parser,
                split_matching,
                output_key=output_key,
                rng=rng,
            )
        if combined and form_index < len(forms) - 1:
            doc.add_page_break()

    doc.save(outfile)


def setup_a4(document):
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(30)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)
    return document


def setup_metadata(document):
    properties = document.core_properties
    properties.author = "qti-converter"
    properties.title = "Quiz export from LMS"
    return document
